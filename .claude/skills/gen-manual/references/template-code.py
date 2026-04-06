"""
gen-manual 스킬 — Word 매뉴얼 생성 기반 코드 패턴
=======================================================
이 파일은 Claude가 매뉴얼 생성 스크립트를 작성할 때 참고하는 패턴 파일입니다.
실제 생성 시에는 상단 DATA 섹션 변수를 프로그램 정보로 채워 사용합니다.

사용법:
  1. 아래 [DATA 섹션]의 변수를 대상 프로그램 정보로 교체
  2. SECTIONS 리스트에서 불필요한 섹션 제거 또는 추가
  3. OUTPUT_PATH를 원하는 저장 경로로 설정
  4. python gen_manual.py 실행
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

# ════════════════════════════════════════════════════════════
# [DATA 섹션] — 여기만 수정하여 사용
# ════════════════════════════════════════════════════════════

APP_NAME        = "프로그램명"
VERSION         = "v1.0"
TARGET_AUDIENCE = "일반 사용자"       # 영어 예: "General Users"
TODAY           = datetime.date.today().strftime("%Y-%m-%d")
OUTPUT_PATH     = r"manuals\YYYYMMDD_프로그램명_매뉴얼.docx"

# 언어 설정: "ko" (한국어, 기본값) 또는 "en" (영어)
LANG = "ko"

# ════════════════════════════════════════════════════════════
# 다국어 레이블 (i18n) — 섹션 제목·표지·목차 등 고정 문자열
# ════════════════════════════════════════════════════════════
I18N = {
    "ko": {
        "subtitle":         "사용자 매뉴얼",
        "cover_version":    "버전",
        "cover_date":       "작성일",
        "cover_audience":   "대상",
        "toc_title":        "목  차",
        "sec_overview":     "개요",
        "sec_prerequisites":"사전 준비사항",
        "sec_getting_started":"시작하기",
        "sec_features":     "주요 기능",
        "sec_data_storage": "데이터 저장 방식",
        "sec_cautions":     "주의사항 및 제한사항",
        "sec_faq":          "문제 해결 (FAQ)",
        "warning_prefix":   "주의",
        "manual_suffix":    "매뉴얼",
    },
    "en": {
        "subtitle":         "User Manual",
        "cover_version":    "Version",
        "cover_date":       "Date",
        "cover_audience":   "Audience",
        "toc_title":        "Table of Contents",
        "sec_overview":     "Overview",
        "sec_prerequisites":"Prerequisites",
        "sec_getting_started":"Getting Started",
        "sec_features":     "Key Features",
        "sec_data_storage": "Data Storage",
        "sec_cautions":     "Cautions & Limitations",
        "sec_faq":          "Troubleshooting (FAQ)",
        "warning_prefix":   "Warning",
        "manual_suffix":    "Manual",
    },
}

def t(key):
    """현재 LANG에 맞는 레이블 반환"""
    return I18N.get(LANG, I18N["ko"])[key]

# 개요 텍스트
OVERVIEW_TEXT = (
    "프로그램 한 줄 설명. 어떤 문제를 해결하며, 어떤 방식으로 동작하는지 2~3문장으로 작성."
)

# 개요 스펙 표 [(항목명, 값), ...]
OVERVIEW_SPECS = [
    ("플랫폼",    "웹 브라우저 / 데스크탑 앱 / CLI"),
    ("데이터 저장", "로컬 파일 / 클라우드 / DB"),
    ("출력 형식",  "Excel / Word / PDF 등"),
    ("인터넷 연결", "필요 / 불필요"),
]

# 사전 준비사항 [문자열 리스트]
PREREQUISITES = [
    "Python 3.8 이상",
    "필요한 패키지: pip install xxx",
    "기타 환경 조건",
]

# 시작하기 단계 [문자열 리스트]
GETTING_STARTED = [
    "1단계: 프로그램 실행 방법",
    "2단계: 초기 설정 방법",
    "3단계: 기본 화면 설명",
]

# 주요 기능 섹션 리스트
# 각 항목: {"num": "4.1", "title": "기능명", "description": "설명",
#            "table_headers": [...], "table_rows": [...], "warning": None or "경고 텍스트"}
FEATURES = [
    {
        "num": "4.1",
        "title": "첫 번째 기능",
        "description": "이 기능이 무엇을 하는지 1~2문장 설명.",
        "table_headers": ["항목", "설명", "예시"],
        "table_rows": [
            ("항목 A", "설명 A", "예시 A"),
            ("항목 B", "설명 B", "예시 B"),
        ],
        "warning": None,
    },
    {
        "num": "4.2",
        "title": "두 번째 기능",
        "description": "이 기능이 무엇을 하는지 1~2문장 설명.",
        "table_headers": ["단계", "동작", "결과"],
        "table_rows": [
            ("1", "버튼 클릭", "파일 저장"),
            ("2", "경로 선택", "저장 완료"),
        ],
        "warning": "주의: 저장 중 프로그램을 닫지 마세요.",
    },
]

# 데이터 저장 방식 설명 [문자열 리스트] — 없으면 빈 리스트 []
DATA_STORAGE_NOTES = [
    "데이터는 로컬 파일(data.json)에 저장됩니다.",
    "백업을 위해 정기적으로 파일을 복사해두세요.",
]

# 주의사항 및 제한사항 [문자열 리스트] — 없으면 빈 리스트 []
CAUTIONS = [
    "삭제한 데이터는 복구되지 않습니다.",
    "파일명에 특수문자 사용 시 오류가 발생할 수 있습니다.",
]

# FAQ [(질문, 답변), ...] — 없으면 빈 리스트 []
FAQS = [
    ("Q. 프로그램이 실행되지 않습니다.",
     "Python 버전을 확인하세요. python --version 명령으로 3.8 이상인지 확인하세요."),
    ("Q. 파일이 저장되지 않습니다.",
     "저장 경로에 쓰기 권한이 있는지 확인하세요."),
]

# ════════════════════════════════════════════════════════════
# 색상 상수 (수정 불필요)
# ════════════════════════════════════════════════════════════

COLOR_PRIMARY   = RGBColor(0x00, 0x70, 0xC0)
COLOR_SECONDARY = RGBColor(0x00, 0x4E, 0x8A)
COLOR_ACCENT    = RGBColor(0xFF, 0x66, 0x00)
COLOR_GRAY      = RGBColor(0xF2, 0xF2, 0xF2)
COLOR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TEXT      = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_ROW_ALT   = RGBColor(0xEA, 0xF4, 0xFF)

# ════════════════════════════════════════════════════════════
# 유틸리티 함수 (수정 불필요)
# ════════════════════════════════════════════════════════════

def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, color="AAAAAA", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'bottom', 'left', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(sz))
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_rule(doc, color="0070C0", thickness=6):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def add_section_title(doc, text, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_after = Pt(2)
    add_rule(doc)
    doc.add_paragraph()


def add_subsection_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECONDARY
    p.paragraph_format.space_after = Pt(6)


def add_body_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_TEXT
    p.paragraph_format.space_after = Pt(8)


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_TEXT
    p.paragraph_format.space_after = Pt(3)


def add_warning(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️  {t('warning_prefix')}: {text}")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = COLOR_ACCENT
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(10)


def add_2col_table(doc, rows_data):
    """2열 키-값 표 (개요 스펙 용)"""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = 'Table Grid'
    for i, (key, val) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = val
        set_cell_bg(row.cells[0], COLOR_GRAY)
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_border(cell)
    doc.add_paragraph()


def add_feature_table(doc, headers, rows):
    """기능 설명 표 (헤더 포함 컬러 표)"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    # 헤더 행
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, COLOR_PRIMARY)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_WHITE
        set_cell_border(cell)
    # 데이터 행
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0:
                set_cell_bg(cell, COLOR_ROW_ALT)
            set_cell_border(cell)
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════
# 문서 생성 메인
# ════════════════════════════════════════════════════════════

doc = Document()

# 페이지 여백
section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)


# ── 표지 ──────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(APP_NAME)
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = COLOR_PRIMARY
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(t("subtitle"))
run.font.size = Pt(20)
run.font.color.rgb = COLOR_SECONDARY
p.paragraph_format.space_after = Pt(4)

# 구분선
p_rule = doc.add_paragraph()
pPr = p_rule._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '16')
bottom.set(qn('w:color'), '0070C0')
pBdr.append(bottom)
pPr.append(pBdr)
p_rule.paragraph_format.space_before = Pt(0)
p_rule.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

for label, value in [(t("cover_version"), VERSION), (t("cover_date"), TODAY), (t("cover_audience"), TARGET_AUDIENCE)]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}: ")
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    r2.font.bold = True
    r2.font.color.rgb = COLOR_TEXT
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()


# ── 목차 ──────────────────────────────────────────────────
# [목차 항목은 실제 섹션 구성에 따라 동적으로 생성]
# 형식: (번호, 제목, 페이지)
# 페이지 번호는 추정값 — Word에서 F9로 갱신 권장

p = doc.add_paragraph()
run = p.add_run(t("toc_title"))
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = COLOR_PRIMARY
p.paragraph_format.space_after = Pt(4)
add_rule(doc, thickness=8)
doc.add_paragraph()

toc_items = [
    ("1.", t("sec_overview"),        "3"),
    ("2.", t("sec_prerequisites"),   "3"),
    ("3.", t("sec_getting_started"), "4"),
    ("4.", t("sec_features"),        "4"),
]
# 기능 소제목 동적 추가
for feat in FEATURES:
    toc_items.append((f"  {feat['num']}", feat['title'], "4"))

extra_sections = []
if DATA_STORAGE_NOTES:
    extra_sections.append(("5.", t("sec_data_storage"), "6"))
if CAUTIONS:
    extra_sections.append(("6.", t("sec_cautions"), "6"))
if FAQS:
    extra_sections.append(("7.", t("sec_faq"), "7"))
toc_items.extend(extra_sections)

for num, title, page in toc_items:
    p = doc.add_paragraph()
    is_sub = num.startswith("  ")
    p.paragraph_format.left_indent = Cm(0.5) if is_sub else Cm(0)
    p.paragraph_format.space_after = Pt(3)
    r_n = p.add_run(f"{num}  ")
    r_n.font.size = Pt(10)
    r_n.font.bold = not is_sub
    r_n.font.color.rgb = COLOR_PRIMARY if not is_sub else COLOR_TEXT
    r_t = p.add_run(title)
    r_t.font.size = Pt(10)
    r_t.font.color.rgb = COLOR_TEXT
    dot_count = max(1, 45 - len(title) - len(num.strip()))
    r_d = p.add_run(" " + "." * dot_count + " ")
    r_d.font.size = Pt(10)
    r_d.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    r_p = p.add_run(page)
    r_p.font.size = Pt(10)
    r_p.font.color.rgb = COLOR_TEXT

doc.add_page_break()


# ── 1. 개요 / Overview ────────────────────────────────────
add_section_title(doc, f"1.  {t('sec_overview')}")
add_body_text(doc, OVERVIEW_TEXT)
if OVERVIEW_SPECS:
    add_2col_table(doc, OVERVIEW_SPECS)


# ── 2. 사전 준비사항 / Prerequisites ──────────────────────
if PREREQUISITES:
    add_section_title(doc, f"2.  {t('sec_prerequisites')}")
    for item in PREREQUISITES:
        add_bullet(doc, item)
    doc.add_paragraph()


# ── 3. 시작하기 / Getting Started ─────────────────────────
if GETTING_STARTED:
    add_section_title(doc, f"3.  {t('sec_getting_started')}")
    for item in GETTING_STARTED:
        add_bullet(doc, item)
    doc.add_paragraph()


# ── 4. 주요 기능 / Key Features ───────────────────────────
add_section_title(doc, f"4.  {t('sec_features')}")

for feat in FEATURES:
    add_subsection_title(doc, f"{feat['num']}  {feat['title']}")
    if feat.get("description"):
        add_body_text(doc, feat["description"])
    if feat.get("table_headers") and feat.get("table_rows"):
        add_feature_table(doc, feat["table_headers"], feat["table_rows"])
    if feat.get("warning"):
        add_warning(doc, feat["warning"])


# ── 5. 데이터 저장 방식 / Data Storage ────────────────────
if DATA_STORAGE_NOTES:
    doc.add_page_break()
    add_section_title(doc, f"5.  {t('sec_data_storage')}")
    for item in DATA_STORAGE_NOTES:
        add_bullet(doc, item)
    doc.add_paragraph()


# ── 6. 주의사항 및 제한사항 / Cautions ────────────────────
if CAUTIONS:
    add_section_title(doc, f"6.  {t('sec_cautions')}")
    for item in CAUTIONS:
        add_bullet(doc, item)
    doc.add_paragraph()


# ── 7. FAQ ────────────────────────────────────────────────
if FAQS:
    doc.add_page_break()
    add_section_title(doc, f"7.  {t('sec_faq')}")
    for q, a in FAQS:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        p.paragraph_format.space_after = Pt(2)

        p = doc.add_paragraph()
        run = p.add_run(a)
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_TEXT
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(10)


# ── 저장 ──────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT_PATH) if os.path.dirname(OUTPUT_PATH) else ".", exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"✅ 매뉴얼 저장 완료: {OUTPUT_PATH}")
